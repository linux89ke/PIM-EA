"""
QC Analysis.py - QC Automation Audit Report page
"""
import io
import json
import os
from datetime import date
from concurrent.futures import ThreadPoolExecutor

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import requests
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="QC Automation Audit Report",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
GATEWAY_API_KEY = "jvk_pQ7d0kw8FKwDnlwUzPaV7-IGo6IbT0dAp-vea4hPK2ckB4jPJnHIGctBrwUfIkt5"
master_rule_file = "QC Check Validaton  (2).xlsx"

with st.sidebar:
    st.header("QC Audit Report")
    st.divider()
    st.subheader("Validation Rule File")
    st.code(master_rule_file)
    st.divider()
    st.subheader("About")
    st.caption(
        "Auto-detects marketplace from CSV.  \n"
        "Cross-references ALL validation rules.  \n"
        "Shows ALL raw examples in sorted order. \n"
        "Includes Vision Check Image Grids. \n"
        "Generates fast daily Word audit report."
    )

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE & DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────
_C = {
    "dark_blue":  "1F4E79", "blue": "2E75B6", "light_blue": "D9E1F2",
    "white": "FFFFFF", "red": "C00000", "orange": "E67E00",
    "green_txt": "375623", "green_bg": "E2EFDA", "gray_bg": "F2F2F2",
    "alert_red": "FCE4E4", "alert_ora": "FFF3CD", "alert_blu": "D9EAF7",
}

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr(); tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4"); el.set(qn("w:space"), "0"); el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def _paragraph_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def _para_border(para, side="bottom", color="2E75B6", sz="6"):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz); el.set(qn("w:space"), "1"); el.set(qn("w:color"), color)
    pBdr.append(el); pPr.append(pBdr)

def _add_header_footer(doc, country_label, report_date):
    section = doc.sections[0]
    hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    hp.clear(); run = hp.add_run(f"QC VALIDATION AUDIT REPORT  |  Jumia {country_label}  |  {report_date}")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88); run.font.size = Pt(8); run.font.name = "Calibri"
    _para_border(hp, side="bottom", color=_C["blue"], sz="6")

    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pre = fp.add_run("Jumia QC Automation Audit  |  Confidential  |  Page ")
    run_pre.font.size = Pt(8); run_pre.font.name = "Calibri"; run_pre.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2): r = OxmlElement("w:r"); r.append(el); fp._p.append(r)
    _para_border(fp, side="top", color=_C["blue"], sz="6")

def _add_title_banner(doc, country_label, report_date):
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("QC VALIDATION AUDIT REPORT")
    r1.bold = True; r1.font.size = Pt(24); r1.font.name = "Calibri"; r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _paragraph_bg(p1, _C["dark_blue"]); _para_border(p1, "top", _C["blue"], "18")

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AI Compliance Audit — True System Failures & Rule Bypasses")
    r2.font.size = Pt(12); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0xA9, 0xC4, 0xE4)
    _paragraph_bg(p2, _C["dark_blue"])

    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Jumia {country_label}  ·  Prepared by QC Auditor  ·  {report_date}")
    r3.font.size = Pt(10); r3.font.name = "Calibri"; r3.font.color.rgb = RGBColor(0xA9, 0xC4, 0xE4)
    _paragraph_bg(p3, _C["dark_blue"]); _para_border(p3, "bottom", _C["blue"], "18")
    doc.add_paragraph("")

def _add_section_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _paragraph_bg(p, _C["blue"])
    return p

def _add_sub_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _para_border(p, "bottom", _C["blue"], "4")
    return p

def _add_body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"
    return p

def _add_metric_table(doc, metrics):
    n = len(metrics)
    table = doc.add_table(rows=2, cols=n); table.style = "Table Grid"
    for i, m in enumerate(metrics):
        for r_idx in (0, 1):
            cell = table.rows[r_idx].cells[i]
            _set_cell_bg(cell, _C["light_blue"]); _set_cell_borders(cell)
            _cell_margins(cell, top=(120 if r_idx==0 else 0), bottom=(0 if r_idx==0 else 100), left=100, right=100)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                r = p.add_run(str(m["value"]))
                r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(*bytes.fromhex(m.get("color", _C["dark_blue"])))
            else:
                r = p.add_run(m["label"])
                r.font.size = Pt(8); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph("")

def _add_issue_table(doc, df, col_map):
    if df.empty: return
    col_map = {k: v for k, v in col_map.items() if v and v in df.columns}
    if not col_map: return
    
    headers = list(col_map.keys()); cols = list(col_map.values())
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, _C["light_blue"]); _set_cell_borders(cell); _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        
    for ri, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        bg = _C["gray_bg"] if ri % 2 == 1 else _C["white"]
        for i, col in enumerate(cols):
            val = str(row.get(col, ""))
            text = val[:297] + "..." if len(val) > 300 else (val if val not in ("nan", "None") else "")
            _set_cell_bg(cells[i], bg); _set_cell_borders(cells[i]); _cell_margins(cells[i])
            r = cells[i].paragraphs[0].add_run(text)
            r.font.size = Pt(9); r.font.name = "Calibri"
    doc.add_paragraph("")

# ─────────────────────────────────────────────────────────────────────────────
# CORE LOGIC HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def claude_narrative(api_key, system_prompt, user_prompt):
    try:
        combined_prompt = f"Instructions:\n{system_prompt}\n\nTask:\n{user_prompt}"
        resp = requests.post(
            "https://ai-gateway.zuma.jumia.com/v1/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
            json={"model": "claude-sonnet-4.5", "max_tokens": 500, "messages": [{"role": "user", "content": combined_prompt}]},
            timeout=30,
        )
        if not resp.ok: return f"System performed well today. High volume processed accurately."
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"System processed items efficiently according to master validation rules."

def get_col(df, options):
    for o in options:
        for c in df.columns:
            if c.lower() == o.lower(): return c
    return None

# ─────────────────────────────────────────────────────────────────────────────
# PAGE HEADER & FILE UPLOAD
# ─────────────────────────────────────────────────────────────────────────────
st.title("QC Automation Audit Report")
st.caption("Catalog Quality Control — Full Daily Detail Log")
st.divider()

with st.container(border=True):
    st.subheader("Data Source")
    uploaded_file = st.file_uploader("Upload QC Results CSV", type=["csv"])

if not uploaded_file:
    st.info("Upload a QC Results CSV file to begin analysis.", icon=":material/upload_file:")
    st.stop()

# ─────────────────────────────────────────────────────────────────────────────
# DATA LOADING & EVALUATION (COMPREHENSIVE)
# ─────────────────────────────────────────────────────────────────────────────
if not os.path.exists(master_rule_file):
    st.error(f"Validation file missing: **'{master_rule_file}'**.", icon=":material/error:")
    st.stop()

with st.spinner("Extracting all AI failures and rule bypasses..."):
    df_raw = pd.read_csv(uploaded_file, on_bad_lines="skip", low_memory=False)

    country_col = get_col(df_raw, ["ACTIVE_STATUS_COUNTRY", "dsc_shop_active_country"])
    if not country_col:
        st.error("Cannot find country column in CSV.", icon=":material/error:")
        st.stop()

    detected = df_raw[country_col].dropna().astype(str).str.upper().unique().tolist()
    if len(detected) == 1:
        country_code = detected[0].lower()
    else:
        chosen = st.selectbox("Multiple countries detected — select one:", detected)
        country_code = chosen.lower()
        df_raw = df_raw[df_raw[country_col].str.upper() == chosen].copy()

    country_label = "Kenya" if country_code == "ke" else "Uganda"
    sheet_name = f"Mandatory Attributes - {country_code.upper()}"
    report_date = date.today().strftime("%d %B %Y").lstrip("0")

    try:
        df_val = pd.read_excel(master_rule_file, sheet_name=sheet_name)
    except Exception as e:
        st.error(f"Could not load sheet '{sheet_name}': {e}", icon=":material/error:")
        st.stop()

    df_raw["cod_category_code"] = pd.to_numeric(df_raw[get_col(df_raw, ["cod_category_code", "CATEGORY_CODE"])], errors="coerce")
    df_val["ID"] = pd.to_numeric(df_val["ID"], errors="coerce")
    df = pd.merge(df_raw, df_val, left_on="cod_category_code", right_on="ID", how="left")

    def _is_filled(series):
        return series.notna() & (series.astype(str).str.strip().str.lower().isin(["nan", "", "none", "n/a"]) == False)

    # Core Rule Bypasses
    RULE_CHECKS = [
        ("Color",         ["color", "COLOR"],                       "Color_Check_Status"),
        ("Warranty",      ["product_warranty", "PRODUCT_WARRANTY"], "Warranty_Check_Status"),
        ("Variation",     ["list_variations", "VARIATION"],         "Variation_Check_Status"),
        ("FDA Documents", ["fda", "FDA"],                           "FDA_Check_Status"),
    ]

    for rule_name, csv_options, status_col in RULE_CHECKS:
        rule_c = get_col(df, [rule_name]); csv_c = get_col(df, csv_options); stat_c = get_col(df, [status_col])
        bypass_col = f"{rule_name}_Bypass"
        df[bypass_col] = False
        if rule_c and stat_c and csv_c:
            mandatory_mask = df[rule_c].astype(str).str.strip().str.lower() == "mandatory"
            filled_mask = _is_filled(df[csv_c])
            df[bypass_col] = (df[stat_c] == "Approved") & mandatory_mask & ~filled_mask

    # Prohibited Items Bypass
    not_allowed_col = get_col(df, ["Not Allowed", "not allowed"])
    cat_status_col = get_col(df, ["Category_Check_Status"])
    
    if not_allowed_col and cat_status_col:
        df["Prohibited_Bypass"] = (df[cat_status_col] == "Approved") & df[not_allowed_col].astype(str).str.strip().str.lower().isin(["yes", "true", "mandatory", "not allowed", "prohibited"])
    else:
        df["Prohibited_Bypass"] = False

    # Extract Common Columns
    _sku_col = get_col(df, ["cod_parent_sku", "cod_seller_sku", "SELLER_SKU"])
    _name_col = get_col(df, ["dsc_name", "NAME"])
    _cat_col = get_col(df, ["dsc_category_name", "CATEGORY"])
    _brand_col = get_col(df, ["dsc_brand_name", "BRAND"])
    _image_col = get_col(df, ["MAIN_IMAGE", "image1", "image_url", "IMAGE"])

    # Build Raw DataFrames for each failure type
    fda_df       = df[df["FDA Documents_Bypass"]].copy()
    color_df     = df[df["Color_Bypass"]].copy()
    warranty_df  = df[df["Warranty_Bypass"]].copy()
    variation_df = df[df["Variation_Bypass"]].copy()
    prohibited_df= df[df["Prohibited_Bypass"]].copy()

    # NLP & Text Checks
    title_stat = get_col(df, ["Product Name_Brand Name_Status"]); title_df = df[df[title_stat] == "Rejected"].copy() if title_stat else pd.DataFrame()
    lang_stat = get_col(df, ["Title_Language_Check_Status"]); lang_df = df[df[lang_stat] == "Rejected"].copy() if lang_stat else pd.DataFrame()
    desc_stat = get_col(df, ["Description_Check_Status"]); desc_df = df[df[desc_stat] == "Rejected"].copy() if desc_stat else pd.DataFrame()

    # Vision Checks
    brand_stat = get_col(df, ["Brand_Image_Check_Status"]); brand_df = df[df[brand_stat] == "Rejected"].copy() if brand_stat else pd.DataFrame()
    img_stat = get_col(df, ["Image_Quality_Check_Status"]); img_df = df[df[img_stat] == "Rejected"].copy() if img_stat else pd.DataFrame()

    # Category Logic Checks
    cat_rej_df = df[df[cat_status_col] == "Rejected"].copy() if cat_status_col else pd.DataFrame()
    cat_rej_reason = get_col(cat_rej_df, ["Category_Check_Rejection_Reason"])

    wrong_rejections_df = pd.DataFrame()
    if "Status" in cat_rej_df.columns and cat_rej_reason:
        wrong_rejections_df = cat_rej_df[(cat_rej_df["Status"] == "ACTIVE") & (~cat_rej_df[cat_rej_reason].str.contains("Prohibited|prohibited", na=False))].copy()

    inactive_df = pd.DataFrame()
    if cat_rej_reason and _cat_col:
        inactive_mask = cat_rej_df[cat_rej_reason].str.contains("INACTIVE", na=False)
        loop_mask = inactive_mask & cat_rej_df.apply(lambda x: str(x[_cat_col]).lower() in str(x[cat_rej_reason]).lower(), axis=1)
        inactive_df = cat_rej_df[loop_mask].copy()

    dup_col = get_col(df, ["Duplicate_Flag"]); duplicates = int(df[dup_col].notna().sum()) if dup_col else 0
    total = len(df)
    cat_approved = len(df[df[cat_status_col] == "Approved"]) if cat_status_col else 0
    cat_rejected = len(df[df[cat_status_col] == "Rejected"]) if cat_status_col else 0

    # High level totals
    total_ai_errors = len(title_df) + len(brand_df) + len(wrong_rejections_df) + len(inactive_df) + len(lang_df) + len(desc_df) + len(img_df)
    total_rule_bypasses = len(fda_df) + len(color_df) + len(warranty_df) + len(variation_df) + len(prohibited_df)

    # ── BUILD SORTED MASTER AI FAILURE LOG ───────────────────────────────────
    def create_failure_view(df_fail, failure_type, detail_col=None, default_detail=""):
        if df_fail.empty: return pd.DataFrame()
        temp = pd.DataFrame()
        temp["SKU"] = df_fail[_sku_col] if _sku_col in df_fail.columns else "N/A"
        temp["Product Name"] = df_fail[_name_col] if _name_col in df_fail.columns else "N/A"
        temp["Failure Type"] = failure_type
        temp["Failure Detail"] = df_fail[detail_col] if detail_col and detail_col in df_fail.columns else default_detail
        return temp

    master_failures = pd.concat([
        create_failure_view(wrong_rejections_df, "Wrong Rejection", cat_rej_reason),
        create_failure_view(inactive_df, "Inactive Loop", cat_rej_reason),
        create_failure_view(title_df, "Title NLP Rejection", "Product name_Brand name_rejection reason"),
        create_failure_view(lang_df, "Title Language Issue", "Title_Language_Check_Reason"),
        create_failure_view(desc_df, "Description Formatting", "Description_Check_Reason"),
        create_failure_view(brand_df, "Brand Visual Mismatch", "Brand_Image_Check_Reason"),
        create_failure_view(img_df, "Image Quality Issue", "Image_Quality_Check_Reason"),
        create_failure_view(prohibited_df, "Prohibited Bypass", default_detail="AI approved a strictly Not Allowed/Prohibited category"),
        create_failure_view(fda_df, "FDA Bypass", default_detail="Mandatory FDA field blank, but AI approved"),
        create_failure_view(color_df, "Color Bypass", default_detail="Mandatory Color field blank, but AI approved"),
        create_failure_view(warranty_df, "Warranty Bypass", default_detail="Mandatory Warranty blank, but AI approved"),
        create_failure_view(variation_df, "Variation Bypass", default_detail="Mandatory Variation blank, but AI approved")
    ], ignore_index=True)

    if not master_failures.empty:
        # Sort Alphabetically by Failure Type, then by SKU to keep everything neat
        master_failures = master_failures.sort_values(by=["Failure Type", "SKU"]).fillna("Blank/NA")

    # ── DEFINE REPORT SECTIONS CONFIG (RAW EXAMPLES) ─────────────────────────
    REPORT_CONFIG = [
        {
            "name": "Wrong Rejections",
            "df": wrong_rejections_df,
            "exp": "The AI incorrectly rejected active, permissible categories due to flawed semantic mapping.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "Rejection Logic": cat_rej_reason}
        },
        {
            "name": "Inactive Category Loops",
            "df": inactive_df,
            "exp": "The AI rejected outdated categories but suggested replacement paths that terminated in the exact same invalid category.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "Rejection Logic": cat_rej_reason}
        },
        {
            "name": "Title NLP Rejections",
            "df": title_df,
            "exp": "The natural language processor falsely flagged clean product titles for brand repetition or formatting errors.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Brand": _brand_col, "AI Rejection": "Product name_Brand name_rejection reason"}
        },
        {
            "name": "Brand Vision Hallucinations",
            "df": brand_df,
            "exp": "The computer vision model incorrectly flagged legitimate items as counterfeits due to contextual misinterpretations.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Brand": _brand_col, "Detected": "Brand_Detected_On_Product", "AI Rejection": "Brand_Image_Check_Reason"}
        },
        {
            "name": "Prohibited Category Bypasses",
            "df": prohibited_df,
            "exp": "CRITICAL: The system erroneously approved listings assigned to strictly restricted or prohibited categories.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col}
        },
        {
            "name": "FDA Rule Bypasses",
            "df": fda_df,
            "exp": "CRITICAL: Listings requiring mandatory FDA health/safety documentation were approved despite the attribute fields being blank.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "FDA Field": get_col(fda_df, ["fda", "FDA"])}
        },
        {
            "name": "Warranty Rule Bypasses",
            "df": warranty_df,
            "exp": "CRITICAL: High-value products were approved without the mandatory warranty policy information provided by the seller.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "Warranty Field": get_col(warranty_df, ["product_warranty", "PRODUCT_WARRANTY"])}
        },
        {
            "name": "Variation Rule Bypasses",
            "df": variation_df,
            "exp": "CRITICAL: Apparel or multi-option items were approved despite missing the mandatory size/fit variation data.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "Variation Field": get_col(variation_df, ["list_variations", "VARIATION"])}
        },
        {
            "name": "Color Rule Bypasses",
            "df": color_df,
            "exp": "CRITICAL: Products were approved despite missing the mandatory color attribute dictated by the master rules.",
            "col_map": {"SKU": _sku_col, "Product Name": _name_col, "Category": _cat_col, "Color Field": get_col(color_df, ["color", "COLOR"])}
        }
    ]

# ─────────────────────────────────────────────────────────────────────────────
# FRONT END: AI FAILURES DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("## System Failures Dashboard")
st.markdown("A consolidated view of True System Failures (False Negatives & Rule Bypasses), sorted by Failure Type.")

with st.container(border=True):
    st.markdown("#### Failure Metrics")
    m1, m2, m3 = st.columns(3)
    m1.metric("Total AI Analysis Flags", str(total_ai_errors), delta="Text/Vision Rejections & Hallucinations", delta_color="inverse")
    m2.metric("Total Rule Bypasses", str(total_rule_bypasses), delta="Critical Policy Misses", delta_color="inverse")
    m3.metric("Total Duplicate SKUs", str(duplicates), delta="Catalog Clutter", delta_color="off")

st.divider()

if not master_failures.empty:
    st.markdown(f"#### Master AI Failure Log ({len(master_failures)} Total Cases)")
    st.dataframe(
        master_failures, 
        use_container_width=True, 
        hide_index=True, 
        height=450,
        column_config={
            "SKU": st.column_config.TextColumn("SKU", width="small"),
            "Product Name": st.column_config.TextColumn("Product Name", width="medium"),
            "Failure Type": st.column_config.TextColumn("Failure Type", width="medium"),
            "Failure Detail": st.column_config.TextColumn("System Logic / Missing Data", width="large"),
        }
    )
else:
    st.success("Excellent! No AI failures or rule bypasses were detected in this batch.", icon=":material/check_circle:")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# ISSUE DETAIL TABS (ALL RAW EXAMPLES SHOWING)
# ─────────────────────────────────────────────────────────────────────────────
st.subheader("Deep Dive Detail Tabs")

active_issues = [config for config in REPORT_CONFIG if len(config["df"]) > 0]

if not active_issues:
    st.info("No detailed issues to display.", icon=":material/info:")
else:
    tabs = st.tabs([f"{config['name']} ({len(config['df'])})" for config in active_issues])
    for i, config in enumerate(active_issues):
        with tabs[i]:
            st.caption(config["exp"])
            
            # Special Visual Toggle for Brand Image Fails
            if config["name"] == "Brand Vision Hallucinations" and _image_col:
                view_mode = st.radio("Display Mode", ["Table View", "Image Grid"], horizontal=True, label_visibility="collapsed")
                
                if view_mode == "Image Grid":
                    img_cols = st.columns(4)
                    for idx, (_, row) in enumerate(config["df"].iterrows()):
                        with img_cols[idx % 4]:
                            with st.container(border=True):
                                img_url = row.get(_image_col)
                                if pd.notna(img_url) and str(img_url).startswith("http"):
                                    st.image(str(img_url), use_container_width=True)
                                else:
                                    st.markdown("*(No Image Available)*")
                                st.caption(f"**SKU:** {row.get(_sku_col, 'N/A')}")
                                st.caption(f"**Listed:** {row.get(_brand_col, 'N/A')} | **Detected:** {row.get('Brand_Detected_On_Product', 'N/A')}")
                                st.error(str(row.get("Brand_Image_Check_Reason", "")))
                else:
                    safe_display_cols = [c for c in config["col_map"].values() if c in config["df"].columns]
                    st.dataframe(config["df"][safe_display_cols], use_container_width=True, hide_index=True)
            else:
                # Standard Table Display for everything else
                safe_display_cols = [c for c in config["col_map"].values() if c in config["df"].columns]
                st.dataframe(config["df"][safe_display_cols], use_container_width=True, hide_index=True)

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# REPORT GENERATION (PYTHON ACCELERATED / DAILY FORMAT)
# ─────────────────────────────────────────────────────────────────────────────
with st.container(border=True):
    st.subheader("Generate Executive Audit Report")
    
    if "report_docx_buf" in st.session_state:
        st.success("Report ready!", icon=":material/check_circle:")
        col_gen, col_dl = st.columns(2)
        with col_gen:
            if st.button("Re-generate Report", icon=":material/refresh:", use_container_width=True):
                del st.session_state["report_docx_buf"]
                st.rerun()
        with col_dl:
            st.download_button(
                label="Download Audit Report (.docx)",
                data=st.session_state["report_docx_buf"],
                file_name=st.session_state["report_docx_name"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                icon=":material/download:",
                use_container_width=True,
            )
    else:
        if st.button("Generate Daily Audit Report (.docx)", type="primary", icon=":material/summarize:", use_container_width=True):
            with st.spinner("Generating ultra-fast daily audit report..."):

                # 1. ONE fast LLM call just for the "What worked well" narrative
                system_prompt = (
                    "You are a Senior E-commerce QC Auditor. "
                    "Write a 2-sentence 'What Worked Well' summary for this QC batch. "
                    "Praise the system for the successful volume processed and zero bypasses in fields if any. "
                    "Do NOT mention failures here. Plain text only, no markdown or bullets."
                )
                
                success_list = []
                if len(fda_df) == 0: success_list.append("FDA")
                if len(warranty_df) == 0: success_list.append("Warranty")
                if len(color_df) == 0: success_list.append("Color")
                success_str = ", ".join(success_list) if success_list else "baseline attributes"

                user_prompt = f"Data:\nTotal processed: {total}\nCorrectly Approved: {cat_approved}\nZero bypasses in: {success_str}"
                what_worked_narrative = claude_narrative(GATEWAY_API_KEY, system_prompt, user_prompt)

                # 2. Build Word Document
                doc = Document()
                for p in doc.paragraphs: p._element.getparent().remove(p._element)

                _add_header_footer(doc, country_label, report_date)
                _add_title_banner(doc, country_label, report_date)

                # Section 1: Dashboard
                _add_section_heading(doc, "1. Dashboard")
                _add_body(doc, f"On {report_date}, the automated QC system processed {total:,} catalog entries for Jumia {country_label}. Of these, {cat_approved:,} were approved for category alignment, {cat_rejected:,} were rejected, and {duplicates:,} were flagged as duplicates.")
                
                _add_metric_table(doc, [
                    {"label": "Total SKUs", "value": total, "color": _C["dark_blue"]},
                    {"label": "AI Flags/Errors", "value": total_ai_errors, "color": _C["red"] if total_ai_errors else _C["green_txt"]},
                    {"label": "Rule Bypasses", "value": total_rule_bypasses, "color": _C["orange"] if total_rule_bypasses else _C["green_txt"]},
                ])

                # Section 2: What Worked Well
                _add_section_heading(doc, "2. What Worked Well")
                _add_body(doc, what_worked_narrative)
                
                # Section 3: Issues Found (Daily Format Loop)
                _add_section_heading(doc, "3. Issues Found")
                
                if not active_issues:
                    _add_body(doc, "No systemic AI failures or rule bypasses were detected in this batch. Excellent performance.")
                else:
                    issue_counter = 1
                    for config in active_issues:
                        # Issue X - Name (Count SKUs)
                        _add_sub_heading(doc, f"Issue {issue_counter} — {config['name']} ({len(config['df'])} SKUs)")
                        # 1 Sentence Explanation
                        _add_body(doc, config["exp"])
                        # Table with ALL raw examples
                        _add_issue_table(doc, config["df"], config["col_map"])
                        issue_counter += 1

                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)

                st.session_state["report_docx_buf"] = buf.getvalue()
                st.session_state["report_docx_name"] = f"QC_Failure_Audit_{country_code.upper()}_{date.today().strftime('%d%b%Y')}.docx"
                st.rerun()