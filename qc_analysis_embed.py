"""
qc_analysis_embed.py
────────────────────
Drop-in module that renders the QC Automation Audit Report *inside* the main
dashboard — no separate Streamlit page required.
"""

from __future__ import annotations

import io
import os
import re
import zipfile
from datetime import date

import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS & DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────
GATEWAY_API_KEY = "jvk_pQ7d0kw8FKwDnlwUzPaV7-IGo6IbT0dAp-vea4hPK2ckB4jPJnHIGctBrwUfIkt5"

_QC_ANALYSIS_PATTERNS = [
    r"qc[\s_-]*analy",
    r"qc[\s_-]*audit",
    r"audit[\s_-]*report",
]

_MASTER_RULE_FILE = "QC Check Validaton  (2).xlsx"

_C = {
    "dark_blue":  "1F4E79", "blue": "2E75B6", "light_blue": "D9E1F2",
    "white": "FFFFFF", "red": "C00000", "orange": "E67E00",
    "green_txt": "375623", "green_bg": "E2EFDA", "gray_bg": "F2F2F2",
}

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr(); tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def _paragraph_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def _para_border(para, side="bottom", color="2E75B6", sz="6"):
    pPr = para._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz); el.set(qn("w:space"), "1"); el.set(qn("w:color"), color)
    pb.append(el); pPr.append(pb)

def _add_header_footer(doc, country_label, report_date):
    section = doc.sections[0]
    hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    hp.clear(); run = hp.add_run(f"QC VALIDATION AUDIT REPORT  |  Jumia {country_label}  |  {report_date}")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88); run.font.size = Pt(8); run.font.name = "Calibri"
    _para_border(hp, "bottom", _C["blue"], "6")

    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = fp.add_run("Jumia QC Automation Audit  |  Confidential  |  Page ")
    r0.font.size = Pt(8); r0.font.name = "Calibri"; r0.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    for tag, txt in (("begin", ""), ("", " PAGE "), ("end", "")):
        el = OxmlElement("w:fldChar" if tag else "w:instrText")
        if tag: el.set(qn("w:fldCharType"), tag)
        else: el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve"); el.text = txt
        r = OxmlElement("w:r"); r.append(el); fp._p.append(r)
    _para_border(fp, "top", _C["blue"], "6")

def _add_title_banner(doc, country_label, report_date):
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("QC VALIDATION AUDIT REPORT")
    r1.bold = True; r1.font.size = Pt(24); r1.font.name = "Calibri"; r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _paragraph_bg(p1, _C["dark_blue"]); _para_border(p1, "top", _C["blue"], "18")

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Jumia {country_label}  ·  {report_date}")
    r2.font.size = Pt(10); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0xA9, 0xC4, 0xE4)
    _paragraph_bg(p2, _C["dark_blue"]); _para_border(p2, "bottom", _C["blue"], "18")
    doc.add_paragraph("")

def _add_section_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); _paragraph_bg(p, _C["blue"])
    return p

def _add_sub_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79); _para_border(p, "bottom", _C["blue"], "4")
    return p

def _add_body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"
    return p

def _add_metric_table(doc, metrics):
    n = len(metrics)
    table = doc.add_table(rows=2, cols=n); table.style = "Table Grid"
    for i, m in enumerate(metrics):
        for r_idx in (0, 1):
            cell = table.rows[r_idx].cells[i]
            _set_cell_bg(cell, _C["light_blue"]); _set_cell_borders(cell)
            _cell_margins(cell, top=(120 if r_idx==0 else 0), bottom=(0 if r_idx==0 else 100), left=100, right=100)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                r = p.add_run(str(m["value"])); r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"
                r.font.color.rgb = RGBColor(*bytes.fromhex(m.get("color", _C["dark_blue"])))
            else:
                r = p.add_run(m["label"]); r.font.size = Pt(8); r.font.name = "Calibri"
                r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph("")

def _add_issue_table(doc, df, col_map):
    if df.empty: return
    col_map = {k: v for k, v in col_map.items() if v and v in df.columns}
    if not col_map: return

    headers = list(col_map.keys()); cols = list(col_map.values())
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, _C["light_blue"]); _set_cell_borders(cell); _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for ri, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        bg = _C["gray_bg"] if ri % 2 == 1 else _C["white"]
        for i, col in enumerate(cols):
            val = str(row.get(col, ""))
            text = (val[:297] + "...") if len(val) > 300 else ("" if val in ("nan", "None") else val)
            _set_cell_bg(cells[i], bg); _set_cell_borders(cells[i]); _cell_margins(cells[i])
            r = cells[i].paragraphs[0].add_run(text)
            r.font.size = Pt(9); r.font.name = "Calibri"
    doc.add_paragraph("")


# ─────────────────────────────────────────────────────────────────────────────
# CORE LOGIC HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _get_col(df: pd.DataFrame, options: list[str]) -> str | None:
    for o in options:
        for c in df.columns:
            if c.lower() == o.lower(): return c
    return None

def _is_filled(series: pd.Series) -> pd.Series:
    return series.notna() & ~series.astype(str).str.strip().str.lower().isin(["nan", "", "none", "n/a"])

def _claude_narrative(system_prompt: str, user_prompt: str) -> str:
    try:
        combined = f"Instructions:\n{system_prompt}\n\nTask:\n{user_prompt}"
        resp = requests.post(
            "https://ai-gateway.zuma.jumia.com/v1/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {GATEWAY_API_KEY}"},
            json={"model": "claude-sonnet-4.5", "max_tokens": 500, "messages": [{"role": "user", "content": combined}]},
            timeout=30,
        )
        if resp.ok: return resp.json()["choices"][0]["message"]["content"]
    except Exception: pass
    return "The QC system processed items efficiently according to master validation rules."

def _find_qc_analysis_in_zip(zip_bytes: bytes) -> pd.DataFrame | None:
    if not zip_bytes: return None
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            for info in zf.infolist():
                fname = os.path.basename(info.filename).lower()
                if any(re.search(pat, fname) for pat in _QC_ANALYSIS_PATTERNS):
                    data = zf.read(info)
                    if fname.endswith(".csv"): return pd.read_csv(io.BytesIO(data), on_bad_lines="skip", low_memory=False)
                    elif fname.endswith((".xlsx", ".xls")): return pd.read_excel(io.BytesIO(data))
    except Exception: pass
    return None

def _load_master_rules(country_code: str) -> pd.DataFrame | None:
    if not os.path.exists(_MASTER_RULE_FILE): return None
    sheet = f"Mandatory Attributes - {country_code.upper()}"
    try: return pd.read_excel(_MASTER_RULE_FILE, sheet_name=sheet)
    except Exception: return None


# ─────────────────────────────────────────────────────────────────────────────
# RICH DISPLAY HELPER — builds the full per-row column list for any issue
# ─────────────────────────────────────────────────────────────────────────────

def _build_display_columns(
    df: pd.DataFrame,
    R: dict,
    extra_issue_cols: dict[str, str],
) -> dict[str, str]:
    """
    Return an ordered label→column mapping that always includes core columns
    plus issue-specific columns passed in extra_issue_cols.
    """
    def _c(*opts): return _get_col(df, list(opts))

    core = {
        "SKU":                  R.get("_sku"),
        "Product Name":         R.get("_name"),
        "Seller":               _c("SELLER_NAME", "dsc_shop_seller_name", "seller_name", "shop_name"),
        "Category":             R.get("_cat"),
        "Full Category Path":   _c("Initial_Category_Path", "category_path", "full_category"),
        "AI Top Suggestion":    _c("Top1_Category", "top1_category"),
        "Match Score":          _c("Category_Match_Score", "match_score"),
        "AI Rejection Reason":  _c(
            "Category_Check_Rejection_Reason",
            "Color_Rejection_Reason",
            "Warranty_Rejection_Reason",
            "Variation_Rejection_Reason",
            "FDA_Rejection_Reason",
        ),
        "Brand":                R.get("_brand"),
    }

    specific = extra_issue_cols

    ordered: dict[str, str] = {}
    seen_cols = set()
    for label, col in {**core, **specific}.items():
        if col and col in df.columns and col not in seen_cols:
            ordered[label] = col
            seen_cols.add(col)

    return ordered


# ─────────────────────────────────────────────────────────────────────────────
# ANALYSIS CORE
# ─────────────────────────────────────────────────────────────────────────────
def _run_analysis(df_raw: pd.DataFrame, country_code: str) -> dict:
    df = df_raw.copy()

    df_val = _load_master_rules(country_code)
    cat_code_col = _get_col(df, ["cod_category_code", "CATEGORY_CODE"])
    if df_val is not None and cat_code_col:
        df[cat_code_col] = pd.to_numeric(df[cat_code_col], errors="coerce")
        df_val["ID"] = pd.to_numeric(df_val["ID"], errors="coerce")
        df = pd.merge(df, df_val, left_on=cat_code_col, right_on="ID", how="left")

    _sku   = _get_col(df, ["cod_parent_sku", "cod_seller_sku", "SELLER_SKU", "PRODUCT_SET_SID"])
    _name  = _get_col(df, ["dsc_name", "NAME"])
    _cat   = _get_col(df, ["dsc_category_name", "CATEGORY"])
    _brand = _get_col(df, ["dsc_brand_name", "BRAND"])
    _img   = _get_col(df, ["MAIN_IMAGE", "image1", "image_url", "IMAGE"])

    # ── Original rule checks ─────────────────────────────────────────────────
    RULE_CHECKS = [
        ("Color",         ["color", "COLOR"],                       "Color_Check_Status",    "Color_Rejection_Reason"),
        ("Warranty",      ["product_warranty", "PRODUCT_WARRANTY"], "Warranty_Check_Status", "Warranty_Rejection_Reason"),
        ("Variation",     ["list_variations", "VARIATION"],         "Variation_Check_Status","Variation_Rejection_Reason"),
        ("FDA Documents", ["fda", "FDA"],                           "FDA_Check_Status",      "FDA_Rejection_Reason"),
    ]

    for rule_name, csv_opts, status_col, reason_col in RULE_CHECKS:
        rule_c = _get_col(df, [rule_name]); csv_c = _get_col(df, csv_opts)
        stat_c = _get_col(df, [status_col])

        bypass_col    = f"{rule_name}_Bypass"
        false_rej_col = f"{rule_name}_False_Rej"
        df[bypass_col]    = False
        df[false_rej_col] = False

        if stat_c and csv_c:
            filled_mask = _is_filled(df[csv_c])
            if rule_c:
                mandatory_mask = df[rule_c].astype(str).str.strip().str.lower() == "mandatory"
                df[bypass_col] = (df[stat_c] == "Approved") & mandatory_mask & ~filled_mask
            df[false_rej_col] = (df[stat_c] == "Rejected") & filled_mask

    not_allowed_col = _get_col(df, ["Not Allowed", "not allowed"])
    cat_status_col  = _get_col(df, ["Category_Check_Status"])
    if not_allowed_col and cat_status_col:
        df["Prohibited_Bypass"] = (
            (df[cat_status_col] == "Approved") &
            df[not_allowed_col].astype(str).str.strip().str.lower().isin(["yes", "true", "mandatory", "not allowed", "prohibited"])
        )
    else:
        df["Prohibited_Bypass"] = False

    fda_df_bp        = df[df["FDA Documents_Bypass"]].copy()
    color_df_bp      = df[df["Color_Bypass"]].copy()
    warranty_df_bp   = df[df["Warranty_Bypass"]].copy()
    variation_df_bp  = df[df["Variation_Bypass"]].copy()
    prohibited_df_bp = df[df["Prohibited_Bypass"]].copy()

    fda_df_fr       = df[df["FDA Documents_False_Rej"]].copy()
    color_df_fr     = df[df["Color_False_Rej"]].copy()
    warranty_df_fr  = df[df["Warranty_False_Rej"]].copy()
    variation_df_fr = df[df["Variation_False_Rej"]].copy()

    title_stat = _get_col(df, ["Product Name_Brand Name_Status"])
    title_df   = df[df[title_stat] == "Rejected"].copy() if title_stat else pd.DataFrame()
    lang_stat  = _get_col(df, ["Title_Language_Check_Status"])
    lang_df    = df[df[lang_stat] == "Rejected"].copy() if lang_stat else pd.DataFrame()
    desc_stat  = _get_col(df, ["Description_Check_Status"])
    desc_df    = df[df[desc_stat] == "Rejected"].copy() if desc_stat else pd.DataFrame()
    brand_stat = _get_col(df, ["Brand_Image_Check_Status"])
    brand_df   = df[df[brand_stat] == "Rejected"].copy() if brand_stat else pd.DataFrame()
    img_stat   = _get_col(df, ["Image_Quality_Check_Status"])
    img_df     = df[df[img_stat] == "Rejected"].copy() if img_stat else pd.DataFrame()

    cat_rej_df     = df[df[cat_status_col] == "Rejected"].copy() if cat_status_col else pd.DataFrame()
    cat_rej_reason = _get_col(cat_rej_df, ["Category_Check_Rejection_Reason"])

    wrong_rej_df = pd.DataFrame()
    inactive_df  = pd.DataFrame()
    if not cat_rej_df.empty and cat_rej_reason:
        if "Status" in cat_rej_df.columns:
            valid_intervention_keywords = "Prohibited|prohibited|maternity|pregnancy|baby|kids|boys|girls|pet"
            wrong_rej_df = cat_rej_df[
                (cat_rej_df["Status"] == "ACTIVE") &
                (~cat_rej_df[cat_rej_reason].astype(str).str.contains(valid_intervention_keywords, case=False, na=False))
            ].copy()
        if _cat:
            inactive_mask = cat_rej_df[cat_rej_reason].astype(str).str.contains("INACTIVE", na=False)
            loop_mask     = inactive_mask & cat_rej_df.apply(
                lambda x: str(x[_cat]).lower() in str(x[cat_rej_reason]).lower(), axis=1
            )
            inactive_df = cat_rej_df[loop_mask].copy()

    dup_col      = _get_col(df, ["Duplicate_Flag"])
    duplicates   = int(df[dup_col].notna().sum()) if dup_col else 0
    total        = len(df)
    cat_approved = len(df[df[cat_status_col] == "Approved"]) if cat_status_col else 0
    cat_rejected = len(df[df[cat_status_col] == "Rejected"]) if cat_status_col else 0

    # ── NEW CHECKS ────────────────────────────────────────────────────────────
    review_check_cols = {
        "Category":     cat_status_col,
        "Color":        _get_col(df, ["Color_Check_Status"]),
        "Product Name": _get_col(df, ["Product Name_Brand Name_Status"]),
        "Brand Image":  _get_col(df, ["Brand_Image_Check_Status"]),
    }
    review_limbo_parts = []
    for check_name, col in review_check_cols.items():
        if col and col in df.columns:
            sub = df[df[col] == "Review"].copy()
            if not sub.empty:
                sub = sub.copy()
                sub["Limbo_Check"] = check_name
                review_limbo_parts.append(sub)
    review_limbo_df = pd.concat(review_limbo_parts, ignore_index=True) if review_limbo_parts else pd.DataFrame()

    dup_approved_df = pd.DataFrame()
    if dup_col and cat_status_col:
        dup_approved_df = df[df[dup_col].notna() & (df[cat_status_col] == "Approved")].copy()

    w_fields = ["PRODUCT_WARRANTY", "WARRANTY_ADDRESS", "WARRANTY_DURATION", "WARRANTY_TYPE"]
    existing_w = [c for c in w_fields if _get_col(df, [c])]
    partial_warranty_df = pd.DataFrame()
    if len(existing_w) > 1:
        w_filled = pd.DataFrame({c: _is_filled(df[_get_col(df, [c])]) for c in existing_w})
        partial_mask = w_filled.any(axis=1) & ~w_filled.all(axis=1)
        partial_warranty_df = df[partial_mask].copy()
        def _warranty_summary(row):
            filled   = [c for c in existing_w if _is_filled(pd.Series([row.get(_get_col(df, [c]), "")]))[0]]
            missing  = [c for c in existing_w if c not in filled]
            return f"Filled: {', '.join(filled) or 'None'} | Blank: {', '.join(missing) or 'None'}"
        if not partial_warranty_df.empty:
            partial_warranty_df["Warranty_Gap_Summary"] = partial_warranty_df.apply(_warranty_summary, axis=1)

    price_col = _get_col(df, ["GLOBAL_PRICE"])
    missing_price_df = pd.DataFrame()
    if price_col:
        price_num = pd.to_numeric(df[price_col], errors="coerce")
        missing_price_df = df[price_num.isna()].copy()
        missing_price_df["Price_Issue"] = "No price set"
        sale_col = _get_col(df, ["GLOBAL_SALE_PRICE"])
        if sale_col:
            sale_num = pd.to_numeric(df[sale_col], errors="coerce")
            inverted = df[(sale_num > price_num) & sale_num.notna() & price_num.notna()].copy()
            inverted["Price_Issue"] = "Sale price exceeds original price"
            missing_price_df = pd.concat([missing_price_df, inverted], ignore_index=True)

    LOW_CONF_THRESHOLD = 0.70
    conf_col = _get_col(df, ["Category_Match_Score"])
    low_conf_df = pd.DataFrame()
    if conf_col and cat_status_col:
        conf_num = pd.to_numeric(df[conf_col], errors="coerce")
        low_conf_df = df[(conf_num < LOW_CONF_THRESHOLD) & (df[cat_status_col] == "Approved")].copy()
        if not low_conf_df.empty:
            low_conf_df["Confidence_Score"] = conf_num[low_conf_df.index].round(3)

    color_norm_col  = _get_col(df, ["Color_AI_Normalized"])
    color_raw_col   = _get_col(df, ["COLOR", "color"])
    color_ghost_df  = pd.DataFrame()
    if color_norm_col and color_raw_col:
        color_ghost_df = df[_is_filled(df[color_norm_col]) & ~_is_filled(df[color_raw_col])].copy()

    type_col = _get_col(df, ["TYPE"])
    type_issue_df = pd.DataFrame()
    if type_col:
        content_only = df[df[type_col] == "CONTENT_ONLY"].copy()
        image_only   = df[df[type_col] == "IMAGE_ONLY"].copy()
        type_parts   = []
        if img_stat and not content_only.empty:
            co_img_rej = content_only[content_only[img_stat] == "Rejected"].copy()
            if not co_img_rej.empty:
                co_img_rej["Type_Issue"] = "CONTENT_ONLY rejected for Image Quality — image check shouldn't apply"
                type_parts.append(co_img_rej)
        if desc_stat and not image_only.empty:
            io_desc_rej = image_only[image_only[desc_stat] == "Rejected"].copy()
            if not io_desc_rej.empty:
                io_desc_rej["Type_Issue"] = "IMAGE_ONLY rejected for Description — description check shouldn't apply"
                type_parts.append(io_desc_rej)
        if lang_stat and not image_only.empty:
            io_lang_rej = image_only[image_only[lang_stat] == "Rejected"].copy()
            if not io_lang_rej.empty:
                io_lang_rej["Type_Issue"] = "IMAGE_ONLY rejected for Title Language — language check shouldn't apply"
                type_parts.append(io_lang_rej)
        type_issue_df = pd.concat(type_parts, ignore_index=True) if type_parts else pd.DataFrame()

    skip_col = _get_col(df, ["QC_Skip_Reason"])
    undoc_skip_df = pd.DataFrame()
    all_status_cols = [cat_status_col, _get_col(df, ["Warranty_Check_Status"]), _get_col(df, ["FDA_Check_Status"]), _get_col(df, ["Color_Check_Status"]), _get_col(df, ["Variation_Check_Status"])]
    valid_status_cols = [c for c in all_status_cols if c]
    if skip_col and valid_status_cols:
        skipped_mask = df[valid_status_cols].apply(lambda col: col.astype(str).str.strip().str.lower() == "skipped").any(axis=1)
        undoc_skip_df = df[skipped_mask & ~_is_filled(df[skip_col])].copy()
        if not undoc_skip_df.empty:
            undoc_skip_df["Skip_Detail"] = "Item skipped by QC engine — no reason documented"

    total_ai_errors = sum(len(x) for x in [title_df, brand_df, wrong_rej_df, inactive_df, lang_df, desc_df, img_df, warranty_df_fr, color_df_fr, variation_df_fr, fda_df_fr])
    total_rule_bypasses = sum(len(x) for x in [fda_df_bp, color_df_bp, warranty_df_bp, variation_df_bp, prohibited_df_bp])
    total_new_flags = sum(len(x) for x in [review_limbo_df, dup_approved_df, partial_warranty_df, missing_price_df, low_conf_df, color_ghost_df, type_issue_df, undoc_skip_df])

    def _fv(df_fail, failure_type, detail_col=None, default_detail=""):
        if df_fail.empty: return pd.DataFrame()
        tmp = pd.DataFrame()
        tmp["SKU"]          = df_fail[_sku].values  if _sku  and _sku  in df_fail.columns else "N/A"
        tmp["Product Name"] = df_fail[_name].values if _name and _name in df_fail.columns else "N/A"
        tmp["Failure Type"] = failure_type
        tmp["Failure Detail"] = df_fail[detail_col].values if detail_col and detail_col in df_fail.columns else default_detail
        return tmp

    master_failures = pd.concat([
        _fv(wrong_rej_df,      "Wrong Rejection",           cat_rej_reason),
        _fv(inactive_df,       "Inactive Loop",             cat_rej_reason),
        _fv(title_df,          "Title NLP Rejection",       "Product name_Brand name_rejection reason"),
        _fv(lang_df,           "Title Language Issue",      "Title_Language_Check_Reason"),
        _fv(desc_df,           "Description Formatting",    "Description_Check_Reason"),
        _fv(brand_df,          "Brand Visual Mismatch",     "Brand_Image_Check_Reason"),
        _fv(img_df,            "Image Quality Issue",       "Image_Quality_Check_Reason"),
        _fv(warranty_df_fr,    "Warranty False Rejection",  _get_col(df, ["Warranty_Rejection_Reason"])),
        _fv(color_df_fr,       "Color False Rejection",     _get_col(df, ["Color_Rejection_Reason"])),
        _fv(variation_df_fr,   "Variation False Rejection", _get_col(df, ["Variation_Rejection_Reason"])),
        _fv(fda_df_fr,         "FDA False Rejection",       _get_col(df, ["FDA_Rejection_Reason"])),
        _fv(prohibited_df_bp,  "Prohibited Bypass",         default_detail="AI approved a Not Allowed/Prohibited category"),
        _fv(fda_df_bp,         "FDA Bypass",                default_detail="Mandatory FDA field blank, but AI approved"),
        _fv(color_df_bp,       "Color Bypass",              default_detail="Mandatory Color field blank, but AI approved"),
        _fv(warranty_df_bp,    "Warranty Bypass",           default_detail="Mandatory Warranty blank, but AI approved"),
        _fv(variation_df_bp,   "Variation Bypass",          default_detail="Mandatory Variation blank, but AI approved"),
        _fv(review_limbo_df,   "Review Limbo",              "Limbo_Check",          "Stuck in Review — never resolved"),
        _fv(dup_approved_df,   "Duplicate Approved",        dup_col,                "Duplicate SKU approved without resolution"),
        _fv(partial_warranty_df,"Partial Warranty",         "Warranty_Gap_Summary", "Some warranty fields filled, others blank"),
        _fv(missing_price_df,  "Missing/Bad Price",         "Price_Issue",          "Price data missing or invalid"),
        _fv(low_conf_df,       "Low Confidence Approval",   "Confidence_Score",     f"Approved with score < {LOW_CONF_THRESHOLD}"),
        _fv(color_ghost_df,    "Color AI Ghost Fill",       color_norm_col,         "AI invented color — seller field was blank"),
        _fv(type_issue_df,     "Wrong Check for Type",      "Type_Issue",           "Check applied to incompatible product type"),
        _fv(undoc_skip_df,     "Undocumented Skip",         "Skip_Detail",          "Skipped with no reason recorded"),
    ], ignore_index=True)

    if not master_failures.empty:
        master_failures = master_failures.sort_values(["Failure Type", "SKU"]).fillna("Blank/NA")

    def get_grouped_df(df_group, group_cols, sku_c, name_c):
        if df_group.empty: return pd.DataFrame()
        valid_groups = [c for c in group_cols if c and c in df_group.columns]
        if not valid_groups: return df_group

        df_clean = df_group.copy()
        for c in valid_groups: df_clean[c] = df_clean[c].fillna("Blank/NA")

        actual_sku  = sku_c  if sku_c  and sku_c  in df_group.columns else valid_groups[0]
        actual_name = name_c if name_c and name_c in df_group.columns else valid_groups[0]

        grp = df_clean.groupby(valid_groups).agg(
            Cases=(actual_sku, "count"),
            Affected_SKUs=(actual_sku, lambda x: ", ".join(x.astype(str).unique())),
            Example_Product=(actual_name, lambda x: str(x.iloc[0])[:80] + ("..." if len(str(x.iloc[0])) > 80 else ""))
        ).reset_index()
        return grp.sort_values(by=[valid_groups[0], "Cases"], ascending=[True, False])

    report_config = [
        {"name": "Warranty False Rejections", "df": warranty_df_fr, "exp": "The AI incorrectly rejected listings for missing Warranty data, even though the seller provided the information.", "group_cols": [_get_col(df, ["Warranty_Rejection_Reason"]), _cat, _get_col(df, ["product_warranty", "PRODUCT_WARRANTY"])], "col_map": {"Rejection Logic": _get_col(df, ["Warranty_Rejection_Reason"]), "Category": _cat, "Provided Warranty": _get_col(df, ["product_warranty", "PRODUCT_WARRANTY"]), "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Color False Rejections", "df": color_df_fr, "exp": "The AI incorrectly rejected listings for missing Color data, even though the seller provided the information.", "group_cols": [_get_col(df, ["Color_Rejection_Reason"]), _cat, _get_col(df, ["color", "COLOR"])], "col_map": {"Rejection Logic": _get_col(df, ["Color_Rejection_Reason"]), "Category": _cat, "Provided Color": _get_col(df, ["color", "COLOR"]), "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Variation False Rejections", "df": variation_df_fr, "exp": "The AI incorrectly rejected listings for missing Variation data, even though the seller provided the information.", "group_cols": [_get_col(df, ["Variation_Rejection_Reason"]), _cat, _get_col(df, ["list_variations", "VARIATION"])], "col_map": {"Rejection Logic": _get_col(df, ["Variation_Rejection_Reason"]), "Category": _cat, "Provided Variation": _get_col(df, ["list_variations", "VARIATION"]), "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "FDA False Rejections", "df": fda_df_fr, "exp": "The AI incorrectly rejected listings for missing FDA documents, even though the seller provided them.", "group_cols": [_get_col(df, ["FDA_Rejection_Reason"]), _cat, _get_col(df, ["fda", "FDA"])], "col_map": {"Rejection Logic": _get_col(df, ["FDA_Rejection_Reason"]), "Category": _cat, "Provided FDA": _get_col(df, ["fda", "FDA"]), "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Wrong Rejections", "df": wrong_rej_df, "exp": "AI incorrectly rejected active, permissible categories.", "group_cols": [cat_rej_reason, _cat], "col_map": {"Category": _cat, "Rejection Logic": cat_rej_reason, "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Inactive Category Loops", "df": inactive_df, "exp": "AI rejected outdated categories but suggested replacement paths that looped back.", "group_cols": [cat_rej_reason, _cat], "col_map": {"Category": _cat, "Rejection Logic": cat_rej_reason, "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Title NLP Rejections", "df": title_df, "exp": "NLP falsely flagged clean product titles.", "group_cols": ["Product name_Brand name_rejection reason", _brand], "col_map": {"Brand": _brand, "AI Rejection": "Product name_Brand name_rejection reason", "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Brand Vision Hallucinations", "df": brand_df, "exp": "Vision model incorrectly flagged legitimate items as counterfeits.", "group_cols": ["Brand_Image_Check_Reason", "Brand_Detected_On_Product", _brand], "col_map": {"Detected": "Brand_Detected_On_Product", "Brand": _brand, "AI Rejection": "Brand_Image_Check_Reason", "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Language & Formatting Rejections", "df": lang_df, "exp": "NLP rejected titles for missing weight/quantities, frequently misidentifying non-consumable items as groceries.", "group_cols": ["Title_Language_Check_Reason", _cat], "col_map": {"Category": _cat, "AI Rejection": "Title_Language_Check_Reason", "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Prohibited Bypasses", "df": prohibited_df_bp, "exp": "CRITICAL: System approved listings in strictly restricted categories.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "FDA Rule Bypasses", "df": fda_df_bp, "exp": "CRITICAL: Listings requiring FDA docs were approved with blank fields.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Warranty Rule Bypasses", "df": warranty_df_bp, "exp": "CRITICAL: High-value products approved without mandatory warranty info.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Variation Rule Bypasses", "df": variation_df_bp, "exp": "CRITICAL: Multi-option items approved despite missing variation data.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Color Rule Bypasses", "df": color_df_bp, "exp": "CRITICAL: Products approved despite missing mandatory color attribute.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Review Limbo", "df": review_limbo_df, "exp": "Items stuck in 'Review' status — neither approved nor rejected. These are unresolved and blocking seller progress.", "group_cols": ["Limbo_Check", _cat], "col_map": {"Check Type": "Limbo_Check", "Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Duplicate Approved", "df": dup_approved_df, "exp": "Items flagged as duplicates were approved anyway. These inflate the catalog and may cause buyer confusion.", "group_cols": [_get_col(df, ["SELLER_NAME", "dsc_shop_seller_name"]), _cat], "col_map": {"Seller": _get_col(df, ["SELLER_NAME", "dsc_shop_seller_name"]), "Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Partial Warranty", "df": partial_warranty_df, "exp": "Some warranty fields are filled but others are blank. Incomplete warranty info causes seller disputes and poor customer experience.", "group_cols": ["Warranty_Gap_Summary", _cat], "col_map": {"Warranty Gap": "Warranty_Gap_Summary", "Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Missing / Bad Price", "df": missing_price_df, "exp": "Listings with no price or a sale price higher than the original. These cannot be purchased and damage marketplace trust.", "group_cols": ["Price_Issue", _cat], "col_map": {"Issue": "Price_Issue", "Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Low Confidence Approvals", "df": low_conf_df, "exp": f"AI approved the category but its own confidence score was below {LOW_CONF_THRESHOLD}. These are high-risk miscategorisations.", "group_cols": [_cat, _get_col(df, ["Top1_Category"])], "col_map": {"Category": _cat, "AI Suggested": _get_col(df, ["Top1_Category"]), "Score": "Confidence_Score", "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Color AI Ghost Fill", "df": color_ghost_df, "exp": "AI populated Color_AI_Normalized but the seller's COLOR field was blank. The AI invented a color value instead of reading it from seller data.", "group_cols": [color_norm_col, _cat], "col_map": {"AI Color": color_norm_col, "Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
        {"name": "Wrong Check for Product Type", "df": type_issue_df, "exp": "A QC check was applied to a product type it doesn't apply to (e.g. image check on CONTENT_ONLY, description check on IMAGE_ONLY).", "group_cols": ["Type_Issue", type_col], "col_map": {"Issue": "Type_Issue", "Product Type": type_col, "Cases": "Cases", "Affected SKUs": "Affected_SKUs"}},
        {"name": "Undocumented Skips", "df": undoc_skip_df, "exp": "Items were skipped by the QC engine but no skip reason was recorded. This is an audit gap — every skip must be documented.", "group_cols": [_cat], "col_map": {"Category": _cat, "Cases": "Cases", "Affected SKUs": "Affected_SKUs", "Example Product": "Example_Product"}},
    ]

    return {
        "df": df,
        "_sku": _sku, "_name": _name, "_cat": _cat, "_brand": _brand, "_img": _img,
        "cat_rej_reason": cat_rej_reason,
        "total": total, "cat_approved": cat_approved, "cat_rejected": cat_rejected,
        "duplicates": duplicates,
        "total_ai_errors": total_ai_errors,
        "total_rule_bypasses": total_rule_bypasses,
        "total_new_flags": total_new_flags,
        "master_failures": master_failures,
        "report_config": report_config,
        "active_issues": [c for c in report_config if len(c["df"]) > 0],
        "fda_df_bp": fda_df_bp, "color_df_bp": color_df_bp, "warranty_df_bp": warranty_df_bp,
        "review_limbo_df": review_limbo_df, "dup_approved_df": dup_approved_df,
        "missing_price_df": missing_price_df, "low_conf_df": low_conf_df,
        "color_ghost_df": color_ghost_df, "undoc_skip_df": undoc_skip_df,
        "get_grouped_df": get_grouped_df,
    }


# ─────────────────────────────────────────────────────────────────────────────
# FAST STATE CALLBACKS
# ─────────────────────────────────────────────────────────────────────────────

def _fast_trash_single_item(sku_str: str):
    """Callback triggered instantly by the Image Grid ❌ button."""
    if "_qca_trash" not in st.session_state:
        st.session_state["_qca_trash"] = set()
    st.session_state["_qca_trash"].add(sku_str)

def _fast_trash_from_table(editor_key: str, original_df: pd.DataFrame, sku_col: str):
    """Callback triggered instantly when a user checks 'Delete' in the data editor."""
    edits = st.session_state[editor_key].get("edited_rows", {})
    skus_to_trash = []
    
    for row_idx, changes in edits.items():
        if changes.get("Delete") is True:
            # Look up the actual SKU(s) from the backend DataFrame
            val = original_df.iloc[int(row_idx)][sku_col]
            if val:
                for s in str(val).split(", "):
                    skus_to_trash.append(s.strip())
                    
    if skus_to_trash:
        if "_qca_trash" not in st.session_state:
            st.session_state["_qca_trash"] = set()
        st.session_state["_qca_trash"].update(skus_to_trash)
        # Clear editor state to reset the checkbox on the next quick-render
        del st.session_state[editor_key]


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def render_qc_analysis_tab() -> None:
    st.divider()

    # Initialize Global Trash State
    if "_qca_trash" not in st.session_state:
        st.session_state["_qca_trash"] = set()

    zip_bytes: bytes | None = st.session_state.get("zip_image_source_bytes")
    df_audit: pd.DataFrame | None = None
    source_label = ""

    if zip_bytes:
        df_audit = _find_qc_analysis_in_zip(zip_bytes)
        if df_audit is not None:
            source_label = ":material/folder_zip: Loaded from ZIP"

    if df_audit is None:
        df_audit = st.session_state.get("zip_qc_results")
        if df_audit is not None and not df_audit.empty:
            source_label = ":material/dataset: Using QC Results data"

    if df_audit is None or (isinstance(df_audit, pd.DataFrame) and df_audit.empty):
        with st.expander(":material/manage_search: QC Automation Audit Report — Waiting for Data", expanded=True):
            st.info("Please upload a QC Results file or run the main ZIP pipeline to view the Audit Dashboard.", icon=":material/hourglass_bottom:")
            if st.button(":material/refresh: Force Check / Refresh", key="_qca_empty_refresh", use_container_width=True):
                st.rerun()
        return

    country_col   = _get_col(df_audit, ["ACTIVE_STATUS_COUNTRY", "dsc_shop_active_country"])
    country_code  = "ke"
    country_label = "Kenya"
    if country_col:
        detected = df_audit[country_col].dropna().astype(str).str.upper().unique().tolist()
        if detected:
            country_code  = detected[0].lower()
            country_label = {"ke": "Kenya", "ug": "Uganda", "ng": "Nigeria", "gh": "Ghana", "ma": "Morocco"}.get(country_code, country_code.upper())

    report_date = date.today().strftime("%d %B %Y").lstrip("0")
    cache_key   = f"_qca_results_{len(df_audit)}"

    def _force_wipe_cache():
        if cache_key in st.session_state: del st.session_state[cache_key]
        if "_qca_docx_buf" in st.session_state: del st.session_state["_qca_docx_buf"]
        st.session_state["_qca_trash"] = set()

    with st.expander(f":material/troubleshoot: QC Automation Audit Report — {country_label}", expanded=False):
        colA, colB = st.columns([4, 1])
        with colA:
            st.caption(source_label)
            st.subheader("System Failures Dashboard")
            st.caption("Consolidated view of AI Hallucinations, Rule Bypasses, and Catalog Integrity Issues.")
        with colB:
            st.button(":material/sync: Force AI Recheck", on_click=_force_wipe_cache, use_container_width=True, key="_qca_force_rerun")

        if cache_key not in st.session_state:
            with st.spinner("Running deep QC rule validation…"):
                st.session_state[cache_key] = _run_analysis(df_audit, country_code)
        R = st.session_state[cache_key]

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("AI Analysis Flags",   str(R.get("total_ai_errors", 0)), delta="Hallucinations / Logic Errors", delta_color="inverse")
        m2.metric("Rule Bypasses",       str(R.get("total_rule_bypasses", 0)), delta="Critical Policy Misses", delta_color="inverse")
        m3.metric("Catalog Integrity",   str(R.get("total_new_flags", 0)), delta="Price / Duplicate / Limbo", delta_color="inverse")
        m4.metric("Duplicate SKUs",      str(R.get("duplicates", 0)), delta="Catalog Clutter", delta_color="off")

        st.divider()

        # Master Failure Log with Trash Filter Applied
        clean_master = R["master_failures"]
        if not clean_master.empty:
            clean_master = clean_master[~clean_master["SKU"].isin(st.session_state["_qca_trash"])]
        
        if not clean_master.empty:
            st.markdown(f"#### :material/manage_search: Master Failure Log ({len(clean_master)} cases)")
            st.dataframe(
                clean_master, use_container_width=True, hide_index=True, height=420,
                column_config={"SKU": st.column_config.TextColumn("SKU", width="small"), "Product Name": st.column_config.TextColumn("Product Name", width="medium"), "Failure Type": st.column_config.TextColumn("Failure Type", width="medium"), "Failure Detail": st.column_config.TextColumn("System Logic / Detail", width="large")},
            )
        else:
            st.success("No failures detected in this batch (or all flagged items trashed).", icon=":material/check_circle:")

        st.divider()

        # Deep Dive Tabs
        active = R["active_issues"]
        if active:
            st.markdown("### :material/view_list: Deep Dive Detail Tabs")
            
            # --- Smart Filters & Trash UI ---
            c_filter, c_trash = st.columns([3, 1])
            with c_filter:
                smart_filter = st.radio("Smart Filters:", ["All View", "Missing Images Only", "Suspected Duplicates"], horizontal=True)
            with c_trash:
                if st.session_state["_qca_trash"]:
                    with st.expander(f":material/delete: Trash ({len(st.session_state['_qca_trash'])})"):
                        if st.button("Undo All Deletes", key="_qca_undo_all", use_container_width=True):
                            st.session_state["_qca_trash"] = set()
                            st.rerun()

            tabs = st.tabs([f"{c['name']}" for c in active])
            
            for i, cfg in enumerate(active):
                with tabs[i]:
                    st.info(f"**Issue Description:** {cfg['exp']}", icon=":material/info:")
                    
                    # Apply Trash & Smart Filters
                    tab_df = cfg["df"].copy()
                    if R["_sku"] in tab_df.columns:
                        tab_df = tab_df[~tab_df[R["_sku"]].isin(st.session_state["_qca_trash"])]
                    
                    if smart_filter == "Missing Images Only" and R["_img"] in tab_df.columns:
                        tab_df = tab_df[~_is_filled(tab_df[R["_img"]])]
                    elif smart_filter == "Suspected Duplicates" and "Duplicate_Flag" in tab_df.columns:
                        tab_df = tab_df[tab_df["Duplicate_Flag"].notna()]

                    if tab_df.empty:
                        st.success("No active cases matching the current filters.", icon=":material/check:")
                        continue

                    if cfg["name"] == "Brand Vision Hallucinations" and R["_img"]:
                        mode = st.radio("Display Mode", ["Interactive Table", "Image Grid"], horizontal=True, label_visibility="collapsed", key=f"_qca_brand_mode_{i}")
                        
                        if mode == "Image Grid":
                            img_cols = st.columns(4)
                            for idx, (_, row) in enumerate(tab_df.iterrows()):
                                with img_cols[idx % 4]:
                                    with st.container(border=True):
                                        c1, c2 = st.columns([8, 2])
                                        with c2:
                                            # Lightning-fast Image Grid Deletion using Callback
                                            st.button(
                                                ":material/close:", 
                                                key=f"del_img_{row.get(R['_sku'])}_{i}",
                                                on_click=_fast_trash_single_item,
                                                args=(row.get(R["_sku"]),)
                                            )
                                        with c1:
                                            url = row.get(R["_img"])
                                            if pd.notna(url) and str(url).startswith("http"):
                                                st.image(str(url), use_container_width=True)
                                            else:
                                                st.markdown("*(No Image)*")
                                        
                                        st.caption(f"**SKU:** {row.get(R['_sku'], 'N/A')}")
                                        st.caption(f"**Listed:** {row.get(R['_brand'], 'N/A')} | **Detected:** {row.get('Brand_Detected_On_Product', 'N/A')}")
                                        st.error(str(row.get("Brand_Image_Check_Reason", "")))
                            continue

                    # UI renderer: show full per-row context instead of grouped rows
                    display_col_map = _build_display_columns(tab_df, R, cfg["col_map"])
                    
                    if display_col_map:
                        st.markdown("##### :material/table_chart: Interactive Table")
                        st.caption("Check the 'Delete' box on any row to instantly remove it from the audit.")
                        
                        editor_key = f"_qca_editor_{i}"
                        target_sku_col = R.get("_sku") or list(display_col_map.values())[0]

                        # Use raw tab_df, reorder by requested columns
                        display_df = tab_df[list(display_col_map.values())].copy()
                        
                        # Rename columns for display using the labels from _build_display_columns
                        display_df = display_df.rename(columns={v: k for k, v in display_col_map.items()})

                        # The target_sku_col might have been renamed to "SKU"
                        rename_sku = next((k for k, v in display_col_map.items() if v == target_sku_col), target_sku_col)

                        # Prepend the Boolean Delete Column
                        display_df.insert(0, "Delete", False)

                        st.data_editor(
                            display_df,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "Delete": st.column_config.CheckboxColumn("Delete", help="Check to delete instantly")
                            },
                            key=editor_key,
                            on_change=_fast_trash_from_table,
                            args=(editor_key, tab_df, target_sku_col) # we pass tab_df back to map the delete action correctly
                        )

        st.divider()

        with st.container(border=True):
            st.subheader("Generate Executive Audit Report")
            docx_key  = "_qca_docx_buf"
            dname_key = "_qca_docx_name"

            should_generate = False

            if docx_key in st.session_state:
                st.success("Report ready!", icon=":material/check_circle:")
                col_gen, col_dl = st.columns(2)
                with col_gen:
                    if st.button("Re-generate Report", icon=":material/refresh:", use_container_width=True, key="_qca_regen"):
                        should_generate = True
                with col_dl:
                    st.download_button(
                        label="Download Audit Report (.docx)",
                        data=st.session_state[docx_key],
                        file_name=st.session_state[dname_key],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        icon=":material/download:",
                        use_container_width=True,
                        key="_qca_dl",
                    )
            else:
                if st.button("Generate Daily Audit Report (.docx)", type="primary", icon=":material/summarize:", use_container_width=True, key="_qca_gen"):
                    should_generate = True

            if should_generate:
                with st.spinner("Generating audit report…"):
                    success_list = []
                    if len(R["fda_df_bp"]) == 0:       success_list.append("FDA Bypass")
                    if len(R["warranty_df_bp"]) == 0:  success_list.append("Warranty Bypass")
                    if len(R["color_df_bp"]) == 0:     success_list.append("Color Bypass")
                    if len(R["missing_price_df"]) == 0: success_list.append("Missing Price")
                    if len(R["low_conf_df"]) == 0:     success_list.append("Low Confidence")
                    if len(R["undoc_skip_df"]) == 0:   success_list.append("Undocumented Skips")
                    success_str = ", ".join(success_list) or "baseline attributes"

                    narrative = _claude_narrative(
                        "You are a Senior E-commerce QC Auditor. Write a 2-sentence 'What Worked Well' summary for this QC batch. Praise the system for the successful volume processed and zero issues in the listed areas. Do NOT mention failures. Plain text only.",
                        f"Data:\nTotal processed: {R['total']}\nCorrectly Approved: {R['cat_approved']}\nZero issues in: {success_str}",
                    )

                    doc = Document()
                    for p in doc.paragraphs: p._element.getparent().remove(p._element)

                    _add_header_footer(doc, country_label, report_date)
                    _add_title_banner(doc, country_label, report_date)

                    _add_section_heading(doc, "1. Dashboard")
                    _add_body(doc, f"On {report_date}, the automated QC system processed {R['total']:,} catalog entries for Jumia {country_label}. Of these, {R['cat_approved']:,} were approved, {R['cat_rejected']:,} were rejected, and {R.get('duplicates', 0):,} were flagged as duplicates.")

                    metrics = [
                        {"label": "Total SKUs", "value": R.get("total", 0), "color": _C["dark_blue"]},
                        {"label": "AI Flags/Errors", "value": R.get("total_ai_errors", 0), "color": _C["red"] if R.get("total_ai_errors", 0) else _C["green_txt"]},
                        {"label": "Rule Bypasses", "value": R.get("total_rule_bypasses", 0), "color": _C["orange"] if R.get("total_rule_bypasses", 0) else _C["green_txt"]},
                        {"label": "Integrity Flags", "value": R.get("total_new_flags", 0), "color": _C["orange"] if R.get("total_new_flags", 0) else _C["green_txt"]},
                    ]
                    mt = doc.add_table(rows=2, cols=len(metrics)); mt.style = "Table Grid"
                    for i, m in enumerate(metrics):
                        for r_idx in (0, 1):
                            cell = mt.rows[r_idx].cells[i]
                            _set_cell_bg(cell, _C["light_blue"]); _set_cell_borders(cell); _cell_margins(cell, top=(120 if r_idx == 0 else 0), bottom=(0 if r_idx == 0 else 100))
                            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if r_idx == 0:
                                r = p.add_run(str(m["value"])); r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(*bytes.fromhex(m.get("color", _C["dark_blue"])))
                            else:
                                r = p.add_run(m["label"]); r.font.size = Pt(8); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
                    doc.add_paragraph("")

                    _add_section_heading(doc, "2. What Worked Well")
                    _add_body(doc, narrative)

                    _add_section_heading(doc, "3. Issues Found")
                    active_for_doc = R["active_issues"]
                    if not active_for_doc:
                        _add_body(doc, "No systemic AI failures, rule bypasses, or integrity issues detected. Excellent performance.")
                    else:
                        for idx, cfg in enumerate(active_for_doc, 1):
                            # Filter docx generation by trash state
                            tab_doc_df = cfg["df"].copy()
                            if R["_sku"] in tab_doc_df.columns:
                                tab_doc_df = tab_doc_df[~tab_doc_df[R["_sku"]].isin(st.session_state["_qca_trash"])]
                            
                            if not tab_doc_df.empty:
                                _add_sub_heading(doc, f"Issue {idx} — {cfg['name']} ({len(tab_doc_df)} SKUs)")
                                _add_body(doc, cfg["exp"])
                                grouped_table = R["get_grouped_df"](tab_doc_df, cfg["group_cols"], R["_sku"], R["_name"])
                                _add_issue_table(doc, grouped_table, cfg["col_map"])

                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    docx_bytes = buf.getvalue()

                    if docx_bytes:
                        st.session_state[docx_key]  = docx_bytes
                        st.session_state[dname_key] = f"QC_Failure_Audit_{country_code.upper()}_{date.today().strftime('%d%b%Y')}.docx"
                        st.rerun()