"""
Generates a Word (.docx) audit report from the unified targeted-audit results
(evaluate_all_checks() output), in the same structure as the reference
"QC Audit Report" template:

  1. Dashboard        — total reviewed / approval rate / rejection rate
  2. Overview          — one-paragraph summary
  3. Issues Found       — one "Issue N" section per (check, reason type) that
     needs attention (False Approval / False Rejection / AI Error / Needs
     Manual Review), each with its own findings table. Image extraction
     failures and duplicates are just another check in this same loop —
     nothing lives in a separate disconnected list.

Pure python-docx, no network calls — built entirely from the DataFrame
produced by evaluate_all_checks().
"""
from io import BytesIO
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from targeted_audit_filters import CHECK_ORDER, CHECK_LABELS

_NAVY = RGBColor(0x1F, 0x4E, 0x78)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# Which verdicts actually get their own "Issue" section in the report —
# True Rejection / Skipped are correct pipeline behaviour, so they're only
# counted in the Overview, not tabled individually.
_REPORTABLE_VERDICTS = ["False Approval", "False Rejection", "Needs Manual Review", "AI Error"]

# Plain-English wording for the report. The old text was the same abstract
# sentence for every check ("re-checking against the file's own rules for this
# category shows the rejection was incorrect"), which told the reader nothing
# about WHAT was wrong. These phrases name the actual thing that was checked,
# so a paragraph reads like: "These 3 products were rejected for missing FDA or
# regulatory documents. But this category doesn't require FDA documents, so
# they should have been approved."

# What the product was rejected for.
_REJECTED_FOR = {
    "category": "because the category looked wrong for the product",
    "color": "because the color was missing or not recognised",
    "warranty": "because the warranty details were missing",
    "variation": "because the variation was missing",
    "fda": "for missing FDA or regulatory documents",
    "title_weight": "because the title was missing a weight or volume",
    "title_english": "because the title was not in English",
    "name_brand": "because the product name and the brand name did not match",
    "image_quality": "because the image quality was too low",
    "brand_image": "because a brand was spotted on the product image",
}

# What the category's own rule actually says — the contradiction.
_NOT_REQUIRED = {
    "category": "the category is a valid fit for this product",
    "color": "this category does not require a color",
    "warranty": "this category does not require warranty details",
    "variation": "this category does not require a variation",
    "fda": "this category does not require FDA documents",
    "title_weight": "this category does not require a weight or volume in the title",
    "title_english": "the title is in English",
    "name_brand": "the name and the brand do match",
    "image_quality": "the image is good enough",
    "brand_image": "there is no brand problem with the image",
}

_IS_REQUIRED = {
    "category": "the category is wrong for this product",
    "color": "this category does require a color, and it is missing",
    "warranty": "this category does require warranty details, and they are missing",
    "variation": "this category does require a variation, and it is missing",
    "fda": "this category does require FDA documents, and they are missing",
    "title_weight": "this category does require a weight or volume in the title, and it is missing",
    "title_english": "the title is not in English",
    "name_brand": "the name and the brand do not match",
    "image_quality": "the image quality is too low",
    "brand_image": "there is a brand on the image that should not be there",
}


# Reason types worded "<thing> Present But Rejected" mean something different
# from the rest: the product was NOT missing the attribute at all — the file's
# own data contradicts its own rejection. Saying "this category does not require
# it" there would be plainly wrong, so these get their own phrasing.
# A value that is present but not recognised is not the same as a missing one.
_INVALID_VALUE = {
    "color": "the color that is filled in is not one the system recognises",
}

_VALUE_PRESENT = {
    "warranty": "the warranty details are actually filled in",
    "title_weight": "the title does contain a weight or volume",
    "color": "the color is actually filled in",
    "variation": "the variation is actually filled in",
    "fda": "the FDA documents are actually there",
}


def _n_products(n: int) -> str:
    """'1 product' / '12 products' — the overview read '1 products' otherwise."""
    return f"{n:,} product" if n == 1 else f"{n:,} products"


def _explain(check_key: str, reason_type: str, verdict: str, n: int) -> str:
    """One plain-English sentence describing what this group of products means."""
    if n == 1:
        subject, verb, pronoun, obj, needs = "This product", "was", "It", "it", "needs"
    else:
        subject, verb, pronoun, obj, needs = f"These {n} products", "were", "They", "them", "need"

    if verdict == "False Rejection":
        why = _REJECTED_FOR.get(check_key, "by this check")
        if "present but rejected" in str(reason_type).lower():
            rule = _VALUE_PRESENT.get(check_key, "the value is actually there in the file")
        else:
            rule = _NOT_REQUIRED.get(check_key)
        if rule:
            return (f"{subject} {verb} rejected {why}. But {rule}, so the rejection was "
                    f"wrong — {pronoun.lower()} should have been approved.")
        return (f"{subject} {verb} rejected {why}, but the file's own data shows the "
                f"rejection was wrong.")

    if verdict == "False Approval":
        rt = str(reason_type).lower()
        if "not recognised" in rt or "invalid" in rt:
            rule = _INVALID_VALUE.get(check_key)
        else:
            rule = _IS_REQUIRED.get(check_key)
        if rule:
            return f"{subject} {verb} approved, but {rule}. {pronoun} should have been rejected."
        return (f"{subject} {verb} approved, but the file's own data shows "
                f"{pronoun.lower()} should have been rejected.")

    if verdict == "Needs Manual Review":
        return (f"{subject} cannot be decided from the file alone. Someone needs to look at "
                f"{obj} — usually to check the photo, or to judge a suggestion the AI was "
                f"unsure about.")

    if verdict == "AI Error":
        return (f"{subject} {verb} never actually checked. The check itself failed — normally "
                f"a timeout or a bad reply from the service — so this is not a problem with "
                f"the {'product' if n == 1 else 'products'}. {pronoun} still {needs} to be "
                f"checked properly.")

    return f"{subject} {needs} attention."


def _shade_cell(cell, hex_color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _clean(val) -> str:
    if val is None:
        return ""
    s = str(val)
    return "" if s.strip().lower() in ("nan", "none") else s


def _content_cols(df: pd.DataFrame, exclude: tuple) -> list:
    """Only columns that actually have data in this slice — since every check
    shares one combined results DataFrame, an unrelated check's column would
    otherwise show up entirely empty (e.g. FDA on a Color-check row)."""
    cols = []
    for c in df.columns:
        if c in exclude:
            continue
        vals = df[c].fillna("").astype(str).str.strip().replace({"nan": "", "None": ""})
        if vals.ne("").any():
            cols.append(c)
    return cols


# Explicit, deliberately minimal column set per check for the printed report —
# SKU / Name / Category always, plus at most one or two fields that are
# actually the subject of that check, plus Detail last. Raw image URLs are
# left out of the Word doc (not actionable in print); they're still shown as
# thumbnails in the app itself.
_DOC_EXTRA_FIELDS = {
    "skip": [],
    "duplicate": [],
    "category": ["Suggested Categories", "Top1 Category", "Category Match Score",
                 "Top1 Score", "Initial Category Path"],
    "color": ["Color", "Color (AI Normalized)", "Color Family"],
    "warranty": ["Warranty", "Warranty Type", "Warranty Duration", "Warranty Address"],
    "variation": ["Variation", "Existing Variation Count"],
    "fda": ["FDA"],
    "title_weight": [],
    "title_english": [],
    "name_brand": ["Brand", "Brand Detected On Product"],
    "image_quality": ["Image Filename"],
    "image_extraction": ["Image"],
    "ai_caption": ["Image"],
    "brand_image": ["Listed Brand", "Brand Detected On Product"],
}


def _doc_columns(check_key: str, df: pd.DataFrame) -> list:
    base = [c for c in ("ProductSetSid", "Product Name", "Seller", "Category") if c in df.columns]
    extra = [c for c in _DOC_EXTRA_FIELDS.get(check_key, []) if c in df.columns]
    # Still drop a whitelisted column if it happens to be empty for this
    # particular slice (e.g. Brand blank on some row) — no point in a
    # column of nothing.
    extra = [c for c in extra if _content_cols(df[[c]], exclude=())]
    tail = [c for c in ("Detail",) if c in df.columns]
    return base + extra + tail


def _add_table(doc: Document, headers: list, rows: list, widths_in: list = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Wide tables (many context columns) get a slightly smaller font so
    # nothing gets squeezed unreadably thin.
    font_size = Pt(9) if len(headers) <= 6 else Pt(7.5)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = font_size
        run.font.color.rgb = _WHITE
        _shade_cell(hdr_cells[i], "1F4E78")

    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(_clean(val))
            run.font.size = font_size

    if widths_in:
        for row in table.rows:
            for cell, w in zip(row.cells, widths_in):
                cell.width = Inches(w)
    return table


# Relative column widths (inches, ~6.5in usable page width) so a UUID-style
# SKU doesn't eat the same space as a long Detail explanation.
_COL_WIDTHS = {
    "ProductSetSid": 0.9,
    "Product Name": 1.3,
    "Category": 0.9,
    "Color": 0.6,
    "Color Family": 0.7,
    "Color (AI Normalized)": 0.8,
    "Warranty": 0.7,
    "Warranty Type": 0.7,
    "Warranty Duration": 0.7,
    "Warranty Address": 0.9,
    "Variation": 0.8,
    "Existing Variation Count": 0.7,
    "FDA": 0.7,
    "Brand": 0.7,
    "Listed Brand": 0.7,
    "Brand Detected On Product": 0.9,
    "Seller": 0.7,
    "Suggested Categories": 1.1,
    "Top1 Category": 0.9,
    "Category Match Score": 0.6,
    "Top1 Score": 0.6,
    "Initial Category Path": 0.9,
    "Image Filename": 0.9,
    "Image": 1.2,
    "Detail": 1.6,
}


def _widths_for(headers: list) -> list:
    raw = [_COL_WIDTHS.get(h, 0.8) for h in headers]
    total = sum(raw)
    target = 6.5
    return [round(w / total * target, 2) for w in raw]


def _dashboard_table(doc: Document, total: int, approval_rate: str, rejection_rate: str):
    dash = doc.add_table(rows=2, cols=3)
    dash.style = "Table Grid"
    dash.alignment = WD_TABLE_ALIGNMENT.CENTER
    big_vals = [f"{total:,}", approval_rate, rejection_rate]
    labels = ["Total Products Reviewed", "Approval Rate", "Rejection Rate"]
    for i in range(3):
        top = dash.rows[0].cells[i]
        top.text = ""
        r = top.paragraphs[0].add_run(big_vals[i])
        r.bold = True
        r.font.size = Pt(20)
        top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        bottom = dash.rows[1].cells[i]
        bottom.text = ""
        r2 = bottom.paragraphs[0].add_run(labels[i])
        r2.font.size = Pt(9)
        bottom.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx_report(
    fr: pd.DataFrame,
    results: pd.DataFrame,
    country_label: str = "",
) -> bytes:
    """Returns the .docx file as raw bytes, ready for a Streamlit download_button."""
    doc = Document()

    # ── Title block ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"QC AUDIT REPORT {country_label.upper()}".strip())
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run(f"Prepared by QC Team | {date.today():%d %B %Y}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    # ── 1. Dashboard ───────────────────────────────────────────────────────
    doc.add_heading("1. Dashboard", level=1)
    total = len(fr) if fr is not None and not fr.empty else 0
    if total and "Status" in fr.columns:
        approved = int((fr["Status"] == "Approved").sum())
        rejected = int((fr["Status"] == "Rejected").sum())
    else:
        approved = 0
        rejected = 0
    approval_rate = f"{approved / total * 100:.1f}%" if total else "N/A"
    rejection_rate = f"{rejected / total * 100:.1f}%" if total else "N/A"
    _dashboard_table(doc, total, approval_rate, rejection_rate)
    doc.add_paragraph()

    # ── 2. Overview ────────────────────────────────────────────────────────
    doc.add_heading("2. Overview", level=1)
    counts = results["Verdict"].value_counts() if results is not None and not results.empty else pd.Series(dtype=int)
    n_false_approvals = int(counts.get("False Approval", 0))
    n_false_rej = int(counts.get("False Rejection", 0))
    n_true_rej = int(counts.get("True Rejection", 0))
    n_manual = int(counts.get("Needs Manual Review", 0))
    n_ai_error = int(counts.get("AI Error", 0))
    n_dup = int(counts.get("Duplicate", 0))
    n_skip = int(counts.get("Skipped", 0))

    # Short sentences, one idea each. This was previously a single ~90-word
    # sentence that a reader had to unpick to find the numbers that matter.
    doc.add_paragraph(
        f"QC looked at {total:,} products in total. It approved {approved:,} of them "
        f"and rejected {rejected:,}."
    )
    doc.add_paragraph(
        f"This audit went back over every one of those decisions and checked them "
        f"against the rules for each product's own category. {n_true_rej:,} of the "
        f"rejections were correct and need no action."
    )
    doc.add_paragraph(
        f"The audit found {_n_products(n_false_rej)} that {'was' if n_false_rej == 1 else 'were'} "
        f"rejected when {'it' if n_false_rej == 1 else 'they'} should have passed, and "
        f"{_n_products(n_false_approvals)} that {'was' if n_false_approvals == 1 else 'were'} "
        f"approved when {'it' if n_false_approvals == 1 else 'they'} should have been rejected. "
        f"These are the ones worth acting on first."
    )
    if n_manual:
        doc.add_paragraph(
            f"Another {_n_products(n_manual)} cannot be settled from the file on their own "
            f"and {'needs' if n_manual == 1 else 'need'} a person to look at {'it' if n_manual == 1 else 'them'}."
        )
    if n_ai_error:
        doc.add_paragraph(
            f"{_n_products(n_ai_error)} could not be checked at all, because the check itself "
            f"failed — a broken connection, a timeout, or an image that could not be opened. "
            f"{'This is not a problem' if n_ai_error == 1 else 'These are not problems'} with the "
            f"{'product' if n_ai_error == 1 else 'products'}, but {'it has' if n_ai_error == 1 else 'they have'} not been checked yet."
        )
    if n_dup or n_skip:
        doc.add_paragraph(
            f"Separately, {_n_products(n_dup)} {'was' if n_dup == 1 else 'were'} flagged as duplicates, "
            f"and {n_skip:,} {'was' if n_skip == 1 else 'were'} set aside before QC even started "
            f"because their category is prohibited, inactive, or a tester."
        )

    # ── 3. Issues Found ────────────────────────────────────────────────────
    doc.add_heading("3. Issues Found", level=1)
    issue_num = 1
    any_issue = False

    if results is not None and not results.empty:
        for check_key in CHECK_ORDER:
            label = CHECK_LABELS[check_key]
            check_slice = results[results["Check"] == check_key]
            if check_slice.empty:
                continue

            for reason_type in check_slice["Reason Type"].unique():
                reason_slice = check_slice[check_slice["Reason Type"] == reason_type]

                for verdict in _REPORTABLE_VERDICTS:
                    if check_key == "category" and verdict != "AI Error" and not str(reason_type).startswith("AI:"):
                        continue
                    sub = reason_slice[reason_slice["Verdict"] == verdict]
                    if sub.empty:
                        continue
                    any_issue = True

                    _case_word = "Case" if len(sub) == 1 else "Cases"
                    doc.add_heading(
                        f"Issue {issue_num} — {label}: {reason_type} ({verdict}, {len(sub)} {_case_word})",
                        level=2,
                    )
                    issue_num += 1
                    doc.add_paragraph(_explain(check_key, reason_type, verdict, len(sub)))
                    # AI Error cases (gateway timeouts, bad responses, image
                    # extraction failures) can number in the hundreds on a
                    # rough run and are rarely individually actionable — the
                    # useful signal is the total count above and a couple of
                    # sample rows, not a full table. Capping keeps the report
                    # a normal size instead of ballooning on a bad-connection day.
                    _AI_ERROR_EXAMPLE_CAP = 2
                    table_rows = sub.head(_AI_ERROR_EXAMPLE_CAP) if verdict == "AI Error" else sub
                    cols = _doc_columns(check_key, table_rows)
                    _add_table(doc, cols, table_rows[cols].values.tolist(), widths_in=_widths_for(cols))
                    if verdict == "AI Error" and len(sub) > _AI_ERROR_EXAMPLE_CAP:
                        doc.add_paragraph(
                            f"...showing {_AI_ERROR_EXAMPLE_CAP} of {len(sub)} example(s). "
                            "The remaining cases follow the same pattern and are omitted here to keep this report a manageable size."
                        )
                    doc.add_paragraph()

    if not any_issue:
        doc.add_paragraph("No false approvals, false rejections, AI errors, or manual-review "
                           "cases were found across any check.")

    # ── Infrastructure failures (image extraction) are now covered
    # automatically above — they flow through as an "image_extraction"
    # check with Verdict "AI Error", same as any other check.

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()